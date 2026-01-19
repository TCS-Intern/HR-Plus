"""Resume parsing service for extracting structured data from PDF and DOCX files.

This module provides comprehensive resume parsing capabilities including:
- Contact information extraction (name, email, phone, LinkedIn)
- Summary/objective extraction
- Work experience parsing with dates
- Education history parsing
- Skills extraction
- Certifications identification
"""

import io
import re
from datetime import datetime
from typing import Any

from pypdf import PdfReader
from docx import Document


class ResumeParser:
    """Service for parsing resumes and extracting structured information."""

    # Common section headers for identifying resume sections
    SECTION_PATTERNS = {
        "summary": [
            r"^(?:professional\s+)?summary$",
            r"^(?:career\s+)?objective$",
            r"^profile$",
            r"^about(?:\s+me)?$",
            r"^overview$",
            r"^executive\s+summary$",
        ],
        "experience": [
            r"^(?:work\s+)?experience$",
            r"^(?:professional\s+)?experience$",
            r"^employment(?:\s+history)?$",
            r"^work\s+history$",
            r"^career\s+history$",
            r"^relevant\s+experience$",
        ],
        "education": [
            r"^education(?:al\s+background)?$",
            r"^academic(?:\s+background)?$",
            r"^qualifications$",
            r"^degrees?$",
        ],
        "skills": [
            r"^(?:technical\s+)?skills$",
            r"^core\s+competenc(?:y|ies)$",
            r"^(?:key\s+)?competenc(?:y|ies)$",
            r"^areas?\s+of\s+expertise$",
            r"^proficienc(?:y|ies)$",
            r"^technologies$",
        ],
        "certifications": [
            r"^certification[s]?$",
            r"^licenses?(?:\s+(?:and|&)\s+certifications?)?$",
            r"^certifications?\s+(?:and|&)\s+licenses?$",
            r"^professional\s+certifications?$",
            r"^credentials?$",
        ],
        "projects": [
            r"^projects?$",
            r"^(?:key\s+)?projects?$",
            r"^portfolio$",
            r"^personal\s+projects?$",
        ],
    }

    # Date patterns for parsing experience and education dates
    DATE_PATTERNS = [
        # "January 2020 - Present", "Jan 2020 - Dec 2021"
        r"((?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)\s+\d{4})\s*[-\u2013\u2014to]+\s*(Present|Current|(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)\s+\d{4})",
        # "2020 - Present", "2018 - 2021"
        r"(\d{4})\s*[-\u2013\u2014to]+\s*(Present|Current|\d{4})",
        # "01/2020 - 12/2021", "1/2020 - present"
        r"(\d{1,2}/\d{4})\s*[-\u2013\u2014to]+\s*(Present|Current|\d{1,2}/\d{4})",
    ]

    def __init__(self):
        """Initialize the resume parser."""
        pass

    def extract_text_from_pdf(self, file_content: bytes) -> str:
        """Extract text from a PDF file.

        Args:
            file_content: PDF file content as bytes.

        Returns:
            Extracted text from the PDF.
        """
        try:
            reader = PdfReader(io.BytesIO(file_content))
            text_parts = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
            return "\n".join(text_parts)
        except Exception as e:
            # Return empty string on parsing errors
            return ""

    def extract_text_from_docx(self, file_content: bytes) -> str:
        """Extract text from a DOCX file.

        Args:
            file_content: DOCX file content as bytes.

        Returns:
            Extracted text from the document.
        """
        try:
            doc = Document(io.BytesIO(file_content))
            text_parts = []

            # Extract from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))

            return "\n".join(text_parts)
        except Exception as e:
            # Return empty string on parsing errors
            return ""

    def detect_format(self, filename: str) -> str:
        """Detect file format from filename.

        Args:
            filename: Name of the file.

        Returns:
            File format: 'pdf', 'docx', or 'unknown'.
        """
        filename_lower = filename.lower()
        if filename_lower.endswith(".pdf"):
            return "pdf"
        elif filename_lower.endswith(".docx"):
            return "docx"
        elif filename_lower.endswith(".doc"):
            return "doc"
        return "unknown"

    def extract_text(self, file_content: bytes, filename: str) -> str:
        """Extract text from a document based on file type.

        Args:
            file_content: File content as bytes.
            filename: Original filename to detect type.

        Returns:
            Extracted text.

        Raises:
            ValueError: If file type is not supported.
        """
        file_format = self.detect_format(filename)
        if file_format == "pdf":
            return self.extract_text_from_pdf(file_content)
        elif file_format == "docx":
            return self.extract_text_from_docx(file_content)
        else:
            raise ValueError(f"Unsupported file type: {filename}")

    def extract_contact_info(self, text: str) -> dict[str, str]:
        """Extract contact information from resume text.

        Args:
            text: Raw text from the resume.

        Returns:
            Dictionary with contact fields: name, email, phone, linkedin.
        """
        contact = {
            "name": "",
            "email": "",
            "phone": "",
            "linkedin": "",
        }

        # Extract email
        email_pattern = r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}"
        email_match = re.search(email_pattern, text)
        if email_match:
            contact["email"] = email_match.group(0).lower()

        # Extract phone number
        phone_patterns = [
            # US format: (123) 456-7890, 123-456-7890, 123.456.7890
            r"\+?1?[-.\s]?\(?[0-9]{3}\)?[-.\s]?[0-9]{3}[-.\s]?[0-9]{4}",
            # International format
            r"\+?[0-9]{1,3}[-.\s]?[0-9]{2,4}[-.\s]?[0-9]{3,4}[-.\s]?[0-9]{3,4}",
        ]
        for pattern in phone_patterns:
            phone_match = re.search(pattern, text)
            if phone_match:
                contact["phone"] = phone_match.group(0).strip()
                break

        # Extract LinkedIn URL
        linkedin_pattern = r"(?:https?://)?(?:www\.)?linkedin\.com/in/[a-zA-Z0-9_-]+"
        linkedin_match = re.search(linkedin_pattern, text, re.IGNORECASE)
        if linkedin_match:
            url = linkedin_match.group(0)
            if not url.startswith("http"):
                url = "https://" + url
            contact["linkedin"] = url

        # Extract name (usually at the top of the resume)
        contact["name"] = self._extract_name(text, contact["email"])

        return contact

    def _extract_name(self, text: str, email: str = "") -> str:
        """Extract candidate name from resume text.

        Args:
            text: Raw text from the resume.
            email: Extracted email (used as hint for name detection).

        Returns:
            Extracted name or empty string.
        """
        lines = text.strip().split("\n")

        # Try to find name from email local part
        if email:
            local_part = email.split("@")[0]
            # Common patterns: john.doe, john_doe, johndoe
            name_parts = re.split(r"[._]", local_part)
            if len(name_parts) >= 2:
                potential_name = " ".join(part.capitalize() for part in name_parts[:2])
                # Verify it's not just numbers
                if not re.match(r"^[\d\s]+$", potential_name):
                    pass  # We'll verify against the text below

        # Look at the first few non-empty lines for the name
        # Names are typically at the very top
        for line in lines[:10]:
            line = line.strip()

            # Skip empty lines
            if not line:
                continue

            # Skip lines that look like contact info
            if re.search(r"@|http|www\.|linkedin|phone|email|address|street|city", line, re.IGNORECASE):
                continue

            # Skip lines that look like section headers
            is_header = False
            for patterns in self.SECTION_PATTERNS.values():
                for pattern in patterns:
                    if re.match(pattern, line.lower()):
                        is_header = True
                        break
                if is_header:
                    break
            if is_header:
                continue

            # Skip very long lines (probably not a name)
            if len(line) > 60:
                continue

            # Skip lines with too many numbers
            if len(re.findall(r"\d", line)) > 4:
                continue

            # Check if line looks like a name (2-5 words, mostly letters)
            words = line.split()
            if 1 <= len(words) <= 5:
                # Check if mostly alphabetic
                alpha_chars = sum(1 for c in line if c.isalpha())
                total_chars = sum(1 for c in line if not c.isspace())
                if total_chars > 0 and alpha_chars / total_chars > 0.7:
                    # Clean up and return
                    return " ".join(word.strip(",.") for word in words)

        return ""

    def extract_summary(self, text: str, sections: dict[str, str]) -> str:
        """Extract professional summary from resume.

        Args:
            text: Raw text from the resume.
            sections: Dictionary of identified sections.

        Returns:
            Summary text or empty string.
        """
        # First check if we identified a summary section
        if "summary" in sections:
            return sections["summary"]

        # Try to extract summary from the beginning of the document
        # (often summaries appear before any section headers)
        lines = text.strip().split("\n")

        # Skip the first few lines (usually name and contact info)
        start_idx = 0
        for i, line in enumerate(lines[:10]):
            line = line.strip().lower()
            # Look for where content starts after contact info
            if re.search(r"@|http|www\.|phone|linkedin|\d{3}[-.\s]\d{3}", line):
                start_idx = i + 1

        # Collect lines until we hit a section header
        summary_lines = []
        for line in lines[start_idx:start_idx + 15]:
            line = line.strip()

            # Check if this is a section header
            is_header = False
            for patterns in self.SECTION_PATTERNS.values():
                for pattern in patterns:
                    if re.match(pattern, line.lower()):
                        is_header = True
                        break
                if is_header:
                    break

            if is_header:
                break

            if line and len(line) > 20:  # Substantial content
                summary_lines.append(line)

        if summary_lines:
            return " ".join(summary_lines)

        return ""

    def extract_experience(self, text: str, sections: dict[str, str]) -> list[dict[str, Any]]:
        """Extract work experience entries from resume.

        Args:
            text: Raw text from the resume.
            sections: Dictionary of identified sections.

        Returns:
            List of experience dictionaries with company, title, dates, description.
        """
        experience = []

        section_text = sections.get("experience", "")
        if not section_text:
            return experience

        lines = section_text.split("\n")

        current_entry: dict[str, Any] = {}
        description_lines: list[str] = []

        for line in lines:
            line = line.strip()
            if not line:
                continue

            # Check if this line contains dates
            date_match = None
            for pattern in self.DATE_PATTERNS:
                date_match = re.search(pattern, line, re.IGNORECASE)
                if date_match:
                    break

            # If we found dates, this might be a new entry
            if date_match:
                # Save previous entry
                if current_entry:
                    if description_lines:
                        current_entry["description"] = "\n".join(description_lines)
                    experience.append(current_entry)
                    description_lines = []

                # Start new entry
                current_entry = {
                    "company": "",
                    "title": "",
                    "start_date": date_match.group(1) if date_match else "",
                    "end_date": date_match.group(2) if date_match and len(date_match.groups()) > 1 else "",
                    "description": "",
                }

                # Extract company/title from the line (before or after dates)
                line_without_dates = re.sub(pattern, "", line).strip()
                if line_without_dates:
                    # Try to split into title and company
                    parts = re.split(r"\s+(?:at|@|-|,|\|)\s+", line_without_dates, maxsplit=1)
                    if len(parts) == 2:
                        current_entry["title"] = parts[0].strip()
                        current_entry["company"] = parts[1].strip()
                    else:
                        # Could be just title or company
                        current_entry["title"] = line_without_dates

            elif current_entry:
                # Check if this looks like a title or company line
                if not current_entry.get("company") and re.search(r"Inc\.|LLC|Ltd|Corp|Company|Technologies|Solutions|Group", line, re.IGNORECASE):
                    current_entry["company"] = line
                elif not current_entry.get("title") and len(line) < 60 and not line.startswith(("-", "*", "+")):
                    current_entry["title"] = line
                else:
                    # It's part of the description
                    description_lines.append(line)

        # Don't forget the last entry
        if current_entry:
            if description_lines:
                current_entry["description"] = "\n".join(description_lines)
            experience.append(current_entry)

        return experience

    def extract_education(self, text: str, sections: dict[str, str]) -> list[dict[str, str]]:
        """Extract education entries from resume.

        Args:
            text: Raw text from the resume.
            sections: Dictionary of identified sections.

        Returns:
            List of education dictionaries with school, degree, dates.
        """
        education = []

        section_text = sections.get("education", "")
        if not section_text:
            return education

        lines = section_text.split("\n")

        # Common degree patterns
        degree_patterns = [
            r"(?:Bachelor|Master|Doctor|Ph\.?D|B\.?S\.?|M\.?S\.?|B\.?A\.?|M\.?A\.?|MBA|M\.?B\.?A\.?|B\.?E\.?|M\.?E\.?)",
            r"(?:Associate|Diploma|Certificate)",
        ]
        degree_pattern = "|".join(degree_patterns)

        current_entry: dict[str, str] = {}
        for line in lines:
            line = line.strip()
            if not line:
                continue

            # Check for degree
            degree_match = re.search(degree_pattern, line, re.IGNORECASE)

            # Check for dates
            date_match = None
            for pattern in self.DATE_PATTERNS:
                date_match = re.search(pattern, line, re.IGNORECASE)
                if date_match:
                    break

            # Also check for year-only patterns common in education
            year_match = re.search(r"\b(19|20)\d{2}\b", line)

            if degree_match or (len(line) < 100 and ("University" in line or "College" in line or "Institute" in line or "School" in line)):
                # Save previous entry if exists
                if current_entry:
                    education.append(current_entry)

                current_entry = {
                    "school": "",
                    "degree": "",
                    "field": "",
                    "start_date": "",
                    "end_date": "",
                }

                if degree_match:
                    current_entry["degree"] = line

                if date_match:
                    current_entry["start_date"] = date_match.group(1)
                    current_entry["end_date"] = date_match.group(2) if len(date_match.groups()) > 1 else ""
                elif year_match:
                    current_entry["end_date"] = year_match.group(0)

                # Try to identify school name
                if "University" in line or "College" in line or "Institute" in line or "School" in line:
                    current_entry["school"] = line

            elif current_entry:
                # Additional info for current entry
                if not current_entry.get("school") and ("University" in line or "College" in line or "Institute" in line or "School" in line):
                    current_entry["school"] = line
                elif not current_entry.get("degree") and degree_match:
                    current_entry["degree"] = line
                elif date_match and not current_entry.get("end_date"):
                    current_entry["start_date"] = date_match.group(1)
                    current_entry["end_date"] = date_match.group(2) if len(date_match.groups()) > 1 else ""

        # Don't forget last entry
        if current_entry:
            education.append(current_entry)

        return education

    def extract_skills(self, text: str, sections: dict[str, str]) -> list[str]:
        """Extract skills from resume.

        Args:
            text: Raw text from the resume.
            sections: Dictionary of identified sections.

        Returns:
            List of skills.
        """
        skills = []

        section_text = sections.get("skills", "")
        if not section_text:
            # Try to extract skills from the entire text
            section_text = text

        # Common skill patterns and delimiters
        # Skills are often comma-separated, bullet-pointed, or on separate lines
        lines = section_text.split("\n")

        for line in lines:
            line = line.strip()

            # Skip empty lines and section headers
            if not line or len(line) < 2:
                continue

            # Check if this is a section header
            is_header = False
            for patterns in self.SECTION_PATTERNS.values():
                for pattern in patterns:
                    if re.match(pattern, line.lower()):
                        is_header = True
                        break
                if is_header:
                    break
            if is_header:
                continue

            # Remove bullet points and list markers
            line = re.sub(r"^[-*+\u2022\u2023\u25E6\u2043\u2219]\s*", "", line)

            # Split by common delimiters
            parts = re.split(r"[,;|/]|\s{2,}", line)

            for part in parts:
                part = part.strip()
                # Filter out items that are too short or too long
                if 2 <= len(part) <= 50:
                    # Filter out common non-skill phrases
                    if not re.match(r"^(and|or|the|a|an|with|using|including)$", part.lower()):
                        skills.append(part)

        # Deduplicate while preserving order
        seen = set()
        unique_skills = []
        for skill in skills:
            skill_lower = skill.lower()
            if skill_lower not in seen:
                seen.add(skill_lower)
                unique_skills.append(skill)

        return unique_skills[:50]  # Limit to 50 skills

    def extract_certifications(self, text: str, sections: dict[str, str]) -> list[str]:
        """Extract certifications from resume.

        Args:
            text: Raw text from the resume.
            sections: Dictionary of identified sections.

        Returns:
            List of certifications.
        """
        certifications = []

        section_text = sections.get("certifications", "")

        if section_text:
            lines = section_text.split("\n")
            for line in lines:
                line = line.strip()
                if line and len(line) > 3:
                    # Remove bullet points
                    line = re.sub(r"^[-*+\u2022\u2023\u25E6\u2043\u2219]\s*", "", line)
                    if line:
                        certifications.append(line)
        else:
            # Look for certification patterns throughout the text
            cert_patterns = [
                r"(?:AWS|Amazon Web Services)\s+(?:Certified|Solutions|Developer|Associate|Professional)[^,\n]*",
                r"(?:Google|GCP)\s+(?:Certified|Cloud)[^,\n]*",
                r"(?:Microsoft|Azure)\s+(?:Certified|MCSE|MCSA|MCP)[^,\n]*",
                r"(?:Cisco)\s+(?:CCNA|CCNP|CCIE)[^,\n]*",
                r"(?:PMP|Project Management Professional)",
                r"(?:CISSP|CISM|CEH|CompTIA\s+\w+)",
                r"(?:Scrum Master|PSM|CSM)",
                r"(?:CPA|CFA|CFP|Series\s+\d+)",
            ]

            for pattern in cert_patterns:
                matches = re.findall(pattern, text, re.IGNORECASE)
                certifications.extend(matches)

        # Deduplicate
        seen = set()
        unique_certs = []
        for cert in certifications:
            cert_lower = cert.lower().strip()
            if cert_lower not in seen and len(cert_lower) > 3:
                seen.add(cert_lower)
                unique_certs.append(cert.strip())

        return unique_certs

    def identify_sections(self, text: str) -> dict[str, str]:
        """Identify and extract sections from resume text.

        Args:
            text: Raw text from the resume.

        Returns:
            Dictionary mapping section names to their content.
        """
        sections = {}
        lines = text.split("\n")

        current_section = None
        current_content = []

        for line in lines:
            line_stripped = line.strip()
            line_lower = line_stripped.lower()

            # Check if this line is a section header
            found_section = None
            for section_name, patterns in self.SECTION_PATTERNS.items():
                for pattern in patterns:
                    if re.match(pattern, line_lower):
                        found_section = section_name
                        break
                if found_section:
                    break

            if found_section:
                # Save previous section
                if current_section and current_content:
                    sections[current_section] = "\n".join(current_content).strip()

                current_section = found_section
                current_content = []
            elif current_section:
                current_content.append(line)

        # Save last section
        if current_section and current_content:
            sections[current_section] = "\n".join(current_content).strip()

        return sections

    def parse(self, file_content: bytes, filename: str) -> dict[str, Any]:
        """Parse a resume and extract all structured information.

        Args:
            file_content: Resume file content as bytes.
            filename: Original filename.

        Returns:
            Dictionary with complete parsed resume data including:
            - status: 'success' or 'error'
            - raw_text: Full extracted text
            - contact: {name, email, phone, linkedin}
            - summary: Professional summary text
            - experience: List of work experiences
            - education: List of education entries
            - skills: List of skills
            - certifications: List of certifications
        """
        try:
            # Extract raw text
            raw_text = self.extract_text(file_content, filename)

            if not raw_text.strip():
                return {
                    "status": "error",
                    "error": "Could not extract text from document",
                    "raw_text": "",
                    "contact": {"name": "", "email": "", "phone": "", "linkedin": ""},
                    "summary": "",
                    "experience": [],
                    "education": [],
                    "skills": [],
                    "certifications": [],
                }

            # Identify sections
            sections = self.identify_sections(raw_text)

            # Extract all components
            contact = self.extract_contact_info(raw_text)
            summary = self.extract_summary(raw_text, sections)
            experience = self.extract_experience(raw_text, sections)
            education = self.extract_education(raw_text, sections)
            skills = self.extract_skills(raw_text, sections)
            certifications = self.extract_certifications(raw_text, sections)

            return {
                "status": "success",
                "raw_text": raw_text,
                "contact": contact,
                "summary": summary,
                "experience": experience,
                "education": education,
                "skills": skills,
                "certifications": certifications,
            }

        except ValueError as e:
            return {
                "status": "error",
                "error": str(e),
                "raw_text": "",
                "contact": {"name": "", "email": "", "phone": "", "linkedin": ""},
                "summary": "",
                "experience": [],
                "education": [],
                "skills": [],
                "certifications": [],
            }
        except Exception as e:
            return {
                "status": "error",
                "error": f"Unexpected error parsing resume: {str(e)}",
                "raw_text": "",
                "contact": {"name": "", "email": "", "phone": "", "linkedin": ""},
                "summary": "",
                "experience": [],
                "education": [],
                "skills": [],
                "certifications": [],
            }


# Singleton instance
resume_parser = ResumeParser()
